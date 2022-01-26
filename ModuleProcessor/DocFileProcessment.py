from ast import In
from pydoc import doc
from docx import Document
from docx.shared import Inches
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH






def CreateDoc(filename):
    document = Document()
    document.save(f'{filename}.docx')



def NewDocumentPage(filename,Heading,ques,Code_pic_path,Output_pic_path,futterl,futterR):

    if Heading==None or Heading=="":Heading=" "
    if ques==None or ques=="":ques=" "
    if Code_pic_path==None or Code_pic_path=="":Code_pic_path=" "
    if Output_pic_path==None or Output_pic_path=="":Output_pic_path=" "
    if futterl==None or futterl=="":futterl=" "
    if futterR==None or futterR=="":futterR=" "



    document = Document(f'{filename}.docx')

    p=document.add_heading(f'Assignment: {Heading}',1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(f'{ques}')

    document.add_paragraph('CODE:')

    try:
        document.add_picture(f'ModuleProcessor//LogFiles//{Code_pic_path}',width=Inches(5.25))
    except:pass
    document.add_paragraph('OUTPUT:')

    try:
        document.add_picture(f'ModuleProcessor//LogFiles//{Output_pic_path}',width=Inches(2.25))
    except:pass
    footer_section = document.sections[0]
    footer = footer_section.footer
    footer_text = footer.paragraphs[0]

    footer_text.text = f"{futterl} \t\t{futterR}"

    document.add_page_break()
    
    document.save(f'{filename}.docx')



    